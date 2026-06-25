import docx
import io
import logging

logger = logging.getLogger(__name__)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        texts = []

        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)

        for table in doc.tables:
            texts.append("--- Table ---")
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    texts.append("\t".join(row_texts))

        return "\n".join(texts)
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        raise
